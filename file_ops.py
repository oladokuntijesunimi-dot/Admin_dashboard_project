"This is a custom tool. We are giving the AI the ability to create a physical file on your hard drive."

from langchain_core.tools import tool
from docx import Document

@tool
def save_report_tool(filename: str, content: str):
    """
    Saves the final research report to a Word document.
    Args:
        filename: The name of the file (e.g., 'bitcoin_analysis.docx')
        content: The full text content to write into the document.
    """
    try:
        doc = Document()
        doc.add_heading('Research Report', 0)
        
        # Split content by newlines and parse markdown
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            # Handle Headers
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            # Handle Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                # Check for bold text in bullets
                clean_line = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                if '**' in clean_line:
                    parts = clean_line.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1: # Odd parts were between ** **
                            run.bold = True
                else:
                    p.add_run(clean_line)
            # Handle Standard Paragraphs
            else:
                p = doc.add_paragraph()
                if '**' in line:
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
                else:
                    p.add_run(line)
                
        doc.save(filename)
        return f"File saved successfully: {filename}"
    except Exception as e:
        return f"Error saving file: {str(e)}"
